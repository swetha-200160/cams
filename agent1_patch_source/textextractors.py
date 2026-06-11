from __future__ import annotations

import csv
import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import Literal, Optional

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader
from PIL import Image

from agent1_patch_source.preprocess import deskew_and_enhance


def io_bytes(b: bytes) -> io.BytesIO:
    return io.BytesIO(b)


@dataclass(frozen=True)
class TextResult:
    text: str
    source: Literal["pdf_text", "ocr", "image_ocr", "docx_text", "xlsx_text", "csv_text"]
    pages_used: int
    filename: Optional[str] = field(default=None)


def _try_image_ocr(img: Image.Image, filename: str | None, lang: str) -> str:
    try:
        import pytesseract  # type: ignore
    except Exception:
        return ""
    try:
        return (pytesseract.image_to_string(deskew_and_enhance(img), lang=lang) or "").strip()
    except Exception:
        return ""


def extract_text_from_pdf(
    pdf_bytes: bytes,
    max_pages: int,
    ocr_lang: str,
    filename: str | None = None,
) -> TextResult:
    reader = PdfReader(io_bytes(pdf_bytes))
    pages = min(len(reader.pages), max_pages)
    text_parts: list[str] = []
    for idx in range(pages):
        try:
            txt = (reader.pages[idx].extract_text() or "").strip()
        except Exception:
            txt = ""
        if txt:
            text_parts.append(f"\n\n--- PAGE {idx+1} (PDF TEXT) ---\n{txt}")
    combined = "\n".join(text_parts).strip()
    return TextResult(text=combined, source="pdf_text", pages_used=pages, filename=filename)


def extract_text_from_image(
    image_bytes: bytes,
    ocr_lang: str,
    filename: str | None = None,
) -> TextResult:
    img = Image.open(io_bytes(image_bytes)).convert("RGB")
    text = _try_image_ocr(img, filename, ocr_lang)
    return TextResult(text=text, source="image_ocr", pages_used=1, filename=filename)


def extract_text_from_docx(docx_bytes: bytes, filename: str | None = None) -> TextResult:
    doc = DocxDocument(io_bytes(docx_bytes))
    parts: list[str] = []
    for p in doc.paragraphs:
        s = (p.text or "").strip()
        if s:
            parts.append(s)
    for t in doc.tables:
        try:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        except Exception:
            pass
    return TextResult(text="\n".join(parts).strip(), source="docx_text", pages_used=1, filename=filename)


def extract_text_from_xlsx(xlsx_bytes: bytes, filename: str | None = None) -> TextResult:
    wb = load_workbook(io_bytes(xlsx_bytes), data_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"\n--- SHEET: {sheet.title} ---")
        max_row = min(sheet.max_row or 0, 250)
        max_col = min(sheet.max_column or 0, 50)
        rows: list[list[str]] = []
        for r in range(1, max_row + 1):
            row_vals: list[str] = []
            for c in range(1, max_col + 1):
                v = sheet.cell(row=r, column=c).value
                row_vals.append("" if v is None else str(v).strip())
            if any(row_vals):
                rows.append(row_vals)
        if not rows:
            continue
        parts.append("COLUMNS: " + " | ".join(rows[0]))
        for row_index, row in enumerate(rows[1:], start=2):
            parts.append(f"ROW {row_index}: " + " | ".join(row))
    return TextResult(text="\n".join(parts).strip(), source="xlsx_text", pages_used=1, filename=filename)


def extract_text_from_csv(csv_bytes: bytes, filename: str | None = None) -> TextResult:
    raw = csv_bytes.decode("utf-8", errors="ignore")
    sample = raw[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except Exception:
        dialect = csv.excel
    reader = csv.reader(io.StringIO(raw), dialect)
    rows: list[list[str]] = []
    for idx, row in enumerate(reader, start=1):
        cleaned = [str(cell).strip() for cell in row]
        if any(cleaned):
            rows.append(cleaned)
        if idx >= 500:
            break
    parts: list[str] = [f"--- CSV: {filename or 'uploaded.csv'} ---"]
    if rows:
        parts.append("COLUMNS: " + " | ".join(rows[0]))
        for row_index, row in enumerate(rows[1:], start=2):
            parts.append(f"ROW {row_index}: " + " | ".join(row))
    return TextResult(text="\n".join(parts).strip(), source="csv_text", pages_used=1, filename=filename)


def extract_text(
    file_bytes: bytes,
    mime_type: str,
    filename: str | None = None,
    max_pages: int = 15,
    ocr_lang: str = "eng",
) -> TextResult:
    mt = mime_type.lower().strip()
    lower_name = (filename or "").lower()
    if mt == "application/pdf":
        return extract_text_from_pdf(file_bytes, max_pages, ocr_lang, filename)
    if mt in {"image/png", "image/jpeg", "image/jpg", "image/tiff", "image/bmp", "image/webp"}:
        return extract_text_from_image(file_bytes, ocr_lang, filename)
    if mt in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }:
        return extract_text_from_docx(file_bytes, filename)
    if mt in {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    }:
        return extract_text_from_xlsx(file_bytes, filename)
    if mt in {"text/csv", "application/csv", "text/plain"} and lower_name.endswith(".csv"):
        return extract_text_from_csv(file_bytes, filename)
    raise ValueError(f"Unsupported MIME type '{mime_type}' for {filename or 'file'}")
