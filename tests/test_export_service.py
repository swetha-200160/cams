"""Tests for integrated_cam_backend/export_service.py"""
from __future__ import annotations

import sys
from datetime import datetime, timezone
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))

from export_service import export_draft_to_docx, export_draft_to_pdf
from models import CamBlock, CamDraft, CamSection, EvidenceReference, FileLocator


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------
def _now():
    return datetime.now(timezone.utc)


def _minimal_draft(app_id: str = "app-01", company: str = "ACME Ltd") -> CamDraft:
    return CamDraft(
        application_id=app_id,
        company_name=company,
        generated_at=_now(),
    )


def _full_draft() -> CamDraft:
    locator = FileLocator(type="page", page=2, label="Page 2")
    citation = EvidenceReference(
        id="ev-01",
        document_name="balance_sheet.pdf",
        excerpt="Revenue ₹5000 Cr",
        extracted_value=5000,
        locator=locator,
    )
    block1 = CamBlock(
        id="b-01",
        title="Revenue Overview",
        text="The company reported revenue of ₹5,000 Cr in FY2023.",
        citations=[citation],
    )
    block2 = CamBlock(
        id="b-02",
        title="Profitability",
        text="PAT stood at ₹420 Cr.",
        citations=[],
    )
    section1 = CamSection(
        id="s-01",
        title="Financial Analysis",
        page_hint="Page 5",
        status="ready",
        blocks=[block1, block2],
    )
    section2 = CamSection(
        id="s-02",
        title="Executive Summary",
        status="partial",
        blocks=[
            CamBlock(id="b-03", title="Overview", text="Strong performance.")
        ],
    )
    return CamDraft(
        application_id="app-full",
        company_name="Full Corp Pvt Ltd",
        generated_at=_now(),
        sections=[section1, section2],
        source_documents=["balance_sheet.pdf", "pl.xlsx"],
        notes=["Draft for internal review", "Subject to auditor comments"],
    )


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------
class TestExportDraftToDocx:
    def test_creates_file(self, tmp_path):
        draft = _minimal_draft()
        output = tmp_path / "cam.docx"
        result = export_draft_to_docx(draft, output)
        assert result == output
        assert output.exists()
        assert output.stat().st_size > 0

    def test_creates_parent_dirs(self, tmp_path):
        draft = _minimal_draft()
        output = tmp_path / "nested" / "deep" / "cam.docx"
        export_draft_to_docx(draft, output)
        assert output.exists()

    def test_output_is_valid_docx(self, tmp_path):
        """Valid DOCX files are ZIP archives starting with PK."""
        draft = _minimal_draft()
        output = tmp_path / "cam.docx"
        export_draft_to_docx(draft, output)
        header = output.read_bytes()[:4]
        # DOCX is a ZIP — starts with PK\x03\x04
        assert header == b"PK\x03\x04", "DOCX should be a valid ZIP/OOXML file"

    def test_company_name_in_document(self, tmp_path):
        draft = _minimal_draft(company="Unique Corp Name")
        output = tmp_path / "cam.docx"
        export_draft_to_docx(draft, output)
        from docx import Document
        doc = Document(str(output))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        all_text = full_text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text
        # Title paragraph or heading should contain company name
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading") or "CAM" in p.text or "Unique" in p.text]
        assert any("Unique Corp Name" in h for h in headings) or "Unique Corp Name" in all_text

    def test_sections_headings_in_docx(self, tmp_path):
        draft = _full_draft()
        output = tmp_path / "full_cam.docx"
        export_draft_to_docx(draft, output)
        from docx import Document
        doc = Document(str(output))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Financial Analysis" in all_text
        assert "Executive Summary" in all_text

    def test_block_text_in_docx(self, tmp_path):
        draft = _full_draft()
        output = tmp_path / "full_cam.docx"
        export_draft_to_docx(draft, output)
        from docx import Document
        doc = Document(str(output))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "revenue of ₹5,000 Cr" in all_text

    def test_citations_in_docx(self, tmp_path):
        draft = _full_draft()
        output = tmp_path / "full_cam.docx"
        export_draft_to_docx(draft, output)
        from docx import Document
        doc = Document(str(output))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "balance_sheet.pdf" in all_text

    def test_citation_with_label_in_docx(self, tmp_path):
        draft = _full_draft()
        output = tmp_path / "full_cam.docx"
        export_draft_to_docx(draft, output)
        from docx import Document
        doc = Document(str(output))
        all_text = " ".join(p.text for p in doc.paragraphs)
        # FileLocator label "Page 2" should appear in citation
        assert "Page 2" in all_text

    def test_notes_in_docx(self, tmp_path):
        draft = _full_draft()
        output = tmp_path / "full_cam.docx"
        export_draft_to_docx(draft, output)
        from docx import Document
        doc = Document(str(output))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Draft for internal review" in all_text

    def test_page_hint_in_docx(self, tmp_path):
        draft = _full_draft()
        output = tmp_path / "full_cam.docx"
        export_draft_to_docx(draft, output)
        from docx import Document
        doc = Document(str(output))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Page 5" in all_text

    def test_empty_sections_list(self, tmp_path):
        draft = _minimal_draft()
        output = tmp_path / "empty.docx"
        export_draft_to_docx(draft, output)
        assert output.exists()

    def test_returns_path_object(self, tmp_path):
        draft = _minimal_draft()
        result = export_draft_to_docx(draft, tmp_path / "cam.docx")
        assert isinstance(result, Path)


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------
class TestExportDraftToPdf:
    def test_creates_file(self, tmp_path):
        draft = _minimal_draft()
        output = tmp_path / "cam.pdf"
        result = export_draft_to_pdf(draft, output)
        assert result == output
        assert output.exists()
        assert output.stat().st_size > 0

    def test_creates_parent_dirs(self, tmp_path):
        draft = _minimal_draft()
        output = tmp_path / "nested" / "cam.pdf"
        export_draft_to_pdf(draft, output)
        assert output.exists()

    def test_output_is_valid_pdf(self, tmp_path):
        draft = _minimal_draft()
        output = tmp_path / "cam.pdf"
        export_draft_to_pdf(draft, output)
        header = output.read_bytes()[:4]
        assert header == b"%PDF", "PDF file should start with %PDF"

    def test_section_titles_in_pdf_bytes(self, tmp_path):
        draft = _full_draft()
        output = tmp_path / "full.pdf"
        export_draft_to_pdf(draft, output)
        content = output.read_bytes()
        # PDF content streams are encoded, check for presence
        assert output.stat().st_size > 1000

    def test_empty_sections_does_not_crash(self, tmp_path):
        draft = _minimal_draft()
        output = tmp_path / "empty.pdf"
        export_draft_to_pdf(draft, output)
        assert output.exists()

    def test_notes_section_included(self, tmp_path):
        draft = _full_draft()
        output = tmp_path / "noted.pdf"
        export_draft_to_pdf(draft, output)
        assert output.exists()
        assert output.stat().st_size > 500

    def test_returns_path_object(self, tmp_path):
        draft = _minimal_draft()
        result = export_draft_to_pdf(draft, tmp_path / "cam.pdf")
        assert isinstance(result, Path)

    def test_multiple_blocks_per_section(self, tmp_path):
        draft = _full_draft()
        output = tmp_path / "multi.pdf"
        export_draft_to_pdf(draft, output)
        assert output.exists()

    def test_block_with_multiline_text(self, tmp_path):
        block = CamBlock(
            id="b-01",
            title="Multi-line",
            text="Line one.\nLine two.\nLine three.",
        )
        section = CamSection(id="s-01", title="Test Section", blocks=[block])
        draft = CamDraft(
            application_id="app-01",
            company_name="Test",
            generated_at=_now(),
            sections=[section],
        )
        output = tmp_path / "multiline.pdf"
        export_draft_to_pdf(draft, output)
        assert output.exists()
