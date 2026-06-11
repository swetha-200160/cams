from __future__ import annotations

import io
from pathlib import Path
from xml.sax.saxutils import escape as _xml_escape

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from models import CamDraft


def _safe(text: str) -> str:
    """Escape XML special chars and replace characters unsupported by ReportLab's default font."""
    text = (text or "")
    text = text.replace("₹", "Rs.")
    text = text.replace("\u2013", "-").replace("\u2014", "-")
    text = text.replace("\u2026", "...")
    return _xml_escape(text, entities={"'": "&#39;", '"': "&quot;"})


def export_draft_to_docx(draft: CamDraft, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(f"Draft CAM - {draft.company_name}", level=0)
    doc.add_paragraph(f"Generated at: {draft.generated_at.isoformat()}")

    for section in draft.sections:
        doc.add_heading(section.title, level=1)
        for block in section.blocks:
            if block.title.strip().lower() != section.title.strip().lower():
                doc.add_heading(block.title, level=2)
            doc.add_paragraph(block.text)
            if block.citations:
                cite_para = doc.add_paragraph("Evidence: ")
                for idx, citation in enumerate(block.citations, start=1):
                    label = citation.document_name
                    if citation.locator and citation.locator.label:
                        label += f" ({citation.locator.label})"
                    cite_para.add_run(label)
                    if idx < len(block.citations):
                        cite_para.add_run("; ")

    if draft.notes:
        doc.add_heading("Draft Notes", level=1)
        for note in draft.notes:
            doc.add_paragraph(note, style="List Bullet")

    doc.save(output_path)
    return output_path


def export_draft_to_docx_bytes(draft: CamDraft) -> bytes:
    """Generate a DOCX in memory and return raw bytes (no filesystem write)."""
    buf = io.BytesIO()
    doc = Document()
    doc.add_heading(f"Draft CAM - {draft.company_name}", level=0)
    doc.add_paragraph(f"Generated at: {draft.generated_at.isoformat()}")

    for section in draft.sections:
        doc.add_heading(section.title, level=1)
        for block in section.blocks:
            if block.title.strip().lower() != section.title.strip().lower():
                doc.add_heading(block.title, level=2)
            doc.add_paragraph(block.text)
            if block.citations:
                cite_para = doc.add_paragraph("Evidence: ")
                for idx, citation in enumerate(block.citations, start=1):
                    label = citation.document_name
                    if citation.locator and citation.locator.label:
                        label += f" ({citation.locator.label})"
                    cite_para.add_run(label)
                    if idx < len(block.citations):
                        cite_para.add_run("; ")

    if draft.notes:
        doc.add_heading("Draft Notes", level=1)
        for note in draft.notes:
            doc.add_paragraph(note, style="List Bullet")

    doc.save(buf)
    return buf.getvalue()


def export_draft_to_pdf(draft: CamDraft, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    title_style = styles['Title']
    heading_style = styles['Heading1']
    subheading_style = styles['Heading2']
    body_style = ParagraphStyle(
        'Body',
        parent=styles['BodyText'],
        leading=14,
        spaceAfter=8,
    )
    italic_style = ParagraphStyle(
        'ItalicBody',
        parent=body_style,
        fontName='Helvetica-Oblique',
    )

    story = [
        Paragraph(_safe(f'Draft CAM - {draft.company_name}'), title_style),
        Paragraph(_safe(f'Generated at: {draft.generated_at.isoformat()}'), body_style),
        Spacer(1, 0.15 * inch),
    ]

    for section in draft.sections:
        story.append(Paragraph(_safe(section.title), heading_style))
        for block in section.blocks:
            if block.title.strip().lower() != section.title.strip().lower():
                story.append(Paragraph(_safe(block.title), subheading_style))
            for paragraph_text in (block.text or '').splitlines() or ['']:
                if paragraph_text.strip():
                    story.append(Paragraph(_safe(paragraph_text), body_style))
            if block.citations:
                evidence_labels = []
                for citation in block.citations:
                    label = citation.document_name
                    if citation.locator and citation.locator.label:
                        label += f' ({citation.locator.label})'
                    evidence_labels.append(label)
                story.append(Paragraph(_safe(f"Evidence: {'; '.join(evidence_labels)}"), body_style))
            story.append(Spacer(1, 0.08 * inch))

    if draft.notes:
        story.append(Paragraph('Draft Notes', heading_style))
        for note in draft.notes:
            story.append(Paragraph(_safe(f'• {note}'), body_style))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    doc.build(story)
    return output_path


def export_draft_to_pdf_bytes(draft: CamDraft) -> bytes:
    """Generate a PDF in memory and return raw bytes (no filesystem write)."""
    buf = io.BytesIO()
    styles = getSampleStyleSheet()
    title_style = styles['Title']
    heading_style = styles['Heading1']
    subheading_style = styles['Heading2']
    body_style = ParagraphStyle(
        'Body',
        parent=styles['BodyText'],
        leading=14,
        spaceAfter=8,
    )

    story = [
        Paragraph(_safe(f'Draft CAM - {draft.company_name}'), title_style),
        Paragraph(_safe(f'Generated at: {draft.generated_at.isoformat()}'), body_style),
        Spacer(1, 0.15 * inch),
    ]

    for section in draft.sections:
        story.append(Paragraph(_safe(section.title), heading_style))
        for block in section.blocks:
            if block.title.strip().lower() != section.title.strip().lower():
                story.append(Paragraph(_safe(block.title), subheading_style))
            for paragraph_text in (block.text or '').splitlines() or ['']:
                if paragraph_text.strip():
                    story.append(Paragraph(_safe(paragraph_text), body_style))
            if block.citations:
                evidence_labels = []
                for citation in block.citations:
                    label = citation.document_name
                    if citation.locator and citation.locator.label:
                        label += f' ({citation.locator.label})'
                    evidence_labels.append(label)
                story.append(Paragraph(_safe(f"Evidence: {'; '.join(evidence_labels)}"), body_style))
            story.append(Spacer(1, 0.08 * inch))

    if draft.notes:
        story.append(Paragraph('Draft Notes', heading_style))
        for note in draft.notes:
            story.append(Paragraph(_safe(f'• {note}'), body_style))

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    doc.build(story)
    return buf.getvalue()
